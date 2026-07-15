"""
doc_generator.py
----------------
Generates a polished Word (.docx) itinerary document from the structured
JSON payload returned by the agent.
"""

from __future__ import annotations

import os
from datetime import datetime
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color: str) -> None:
    """Fill a table cell background with a hex colour (e.g. '1F4E79')."""
    tc    = cell._tc
    tcPr  = tc.get_or_add_tcPr()
    shd   = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _add_horizontal_rule(doc: Document) -> None:
    """Insert a thin horizontal rule paragraph."""
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E74B5")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(0)


def _add_colored_heading(doc: Document, text: str,
                         level: int = 1, rgb: tuple = (31, 78, 121)) -> None:
    """Add a heading with a custom RGB colour."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*rgb)
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after  = Pt(4)


def _bullet(doc: Document, text: str, indent: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    p.paragraph_format.left_indent = Inches(0.25 * (indent + 1))
    p.paragraph_format.space_after = Pt(2)


def _kv_row(doc: Document, key: str, value: str,
            key_bold: bool = True) -> None:
    """Add a single label: value paragraph."""
    p    = doc.add_paragraph()
    run1 = p.add_run(f"{key}: ")
    run1.bold = key_bold
    run1.font.color.rgb = RGBColor(31, 78, 121)
    p.add_run(str(value))
    p.paragraph_format.space_after = Pt(2)


def _inr(value: Any, fallback: str = "—") -> str:
    """Safely format a value as an INR amount with thousands separator.
    Handles int, float, numeric strings, and None without crashing."""
    if value is None or value == "" or value == "—":
        return fallback
    try:
        return f"{int(float(str(value).replace(',', ''))):,}"
    except (ValueError, TypeError):
        return str(value)


# ---------------------------------------------------------------------------
# Main document generator
# ---------------------------------------------------------------------------

def generate_itinerary_doc(itinerary: dict[str, Any], user_id: str,
                            output_dir: str = ".") -> str:
    """
    Convert a structured itinerary dict into a formatted .docx file.
    Returns the absolute path to the generated file.
    """
    doc = Document()

    # ---- Page margins ----
    for section in doc.sections:
        section.top_margin    = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # ======================================================
    # COVER SECTION
    # ======================================================
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run  = title_para.add_run("✈  AI Travel Itinerary")
    title_run.font.size   = Pt(28)
    title_run.font.bold   = True
    title_run.font.color.rgb = RGBColor(31, 78, 121)

    dest = itinerary.get("destination", "Your Destination")
    sub_para      = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run       = sub_para.add_run(str(dest).upper())
    sub_run.font.size   = Pt(16)
    sub_run.font.color.rgb = RGBColor(68, 114, 196)
    sub_run.font.bold  = True

    meta_para      = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_para.add_run(
        f"Generated on {datetime.now().strftime('%d %B %Y')}  |  User: {user_id}"
    ).font.color.rgb = RGBColor(128, 128, 128)

    _add_horizontal_rule(doc)
    doc.add_paragraph()

    # ======================================================
    # TRIP OVERVIEW TABLE
    # ======================================================
    _add_colored_heading(doc, "📋 Trip Overview", level=2)

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    _set_cell_bg(hdr[0], "1F4E79")
    _set_cell_bg(hdr[1], "1F4E79")
    for cell, text in zip(hdr, ["Detail", "Information"]):
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)

    overview_fields = [
        ("Destination",      itinerary.get("destination", "—")),
        ("Duration",         f"{itinerary.get('duration_days', '—')} days"),
        ("Group Size",       f"{itinerary.get('group_size', '—')} travellers"),
        ("Total Budget",     f"Rs.{_inr(itinerary.get('budget_total_inr'))}"),
        ("Estimated Cost",   f"Rs.{_inr(itinerary.get('total_estimated_cost_inr'))}"),
        ("Budget Status",    itinerary.get("budget_fit", "—")),
    ]

    for i, (k, v) in enumerate(overview_fields):
        row   = tbl.add_row()
        bg    = "D6E4F0" if i % 2 == 0 else "EBF5FB"
        _set_cell_bg(row.cells[0], bg)
        _set_cell_bg(row.cells[1], bg)
        row.cells[0].paragraphs[0].add_run(k).bold = True
        row.cells[1].paragraphs[0].add_run(str(v))

    doc.add_paragraph()

    # ======================================================
    # SUMMARY
    # ======================================================
    _add_colored_heading(doc, "🗺️  Summary", level=2)
    summary_para = doc.add_paragraph(itinerary.get("summary", ""))
    summary_para.paragraph_format.space_after = Pt(8)

    # ======================================================
    # WEATHER INFO
    # ======================================================
    weather = itinerary.get("weather_info", {})
    if weather:
        _add_colored_heading(doc, "🌤️  Weather at Destination", level=2)
        _kv_row(doc, "Condition",  weather.get("condition", "—"))
        _kv_row(doc, "Temperature",
                f"{weather.get('temp_c', '—')}°C / {weather.get('temp_f', '—')}°F")
        _kv_row(doc, "Best Season", weather.get("best_time", "—"))
        doc.add_paragraph()

    # ======================================================
    # FLIGHT INFO
    # ======================================================
    flight = itinerary.get("flight_info", {})
    if flight:
        _add_colored_heading(doc, "✈️  Flight Information", level=2)
        _kv_row(doc, "Route",     f"{flight.get('origin', '—')} → {flight.get('destination', '—')}")
        _kv_row(doc, "Airline",   flight.get("airline", "—"))
        _kv_row(doc, "Duration",  f"{flight.get('duration_hrs', '—')} hrs")
        _kv_row(doc, "Economy",   f"Rs.{_inr(flight.get('economy'))} per person")
        _kv_row(doc, "Business",  f"Rs.{_inr(flight.get('business'))} per person")
        doc.add_paragraph()

    # ======================================================
    # HOTEL INFO
    # ======================================================
    hotel_info = itinerary.get("hotel_info", {})
    if hotel_info:
        _add_colored_heading(doc, "🏨  Hotel Options", level=2)
        _kv_row(doc, "Tier", hotel_info.get("tier", "—"))
        for option in hotel_info.get("options", []):
            p      = doc.add_paragraph()
            p.add_run(f"  • {option.get('name', '')}").bold = True
            doc.add_paragraph(
                f"    Price: Rs.{_inr(option.get('price_per_night'))}/night  "
                f"|  Rating: {option.get('rating', '—')}  "
                f"|  Amenities: {', '.join(option.get('amenities', []))}"
            ).paragraph_format.left_indent = Inches(0.25)
        doc.add_paragraph()

    # ======================================================
    # DAY-BY-DAY ITINERARY
    # ======================================================
    days = itinerary.get("day_by_day", [])
    if days:
        _add_colored_heading(doc, "📅  Day-by-Day Itinerary", level=2)

        for day_data in days:
            day_num   = day_data.get("day", "?")
            day_title = day_data.get("title", f"Day {day_num}")
            day_cost  = day_data.get("estimated_cost_inr", 0)

            # Day header
            day_para  = doc.add_paragraph()
            day_run   = day_para.add_run(f"Day {day_num}:  {day_title}")
            day_run.bold      = True
            day_run.font.size = Pt(12)
            day_run.font.color.rgb = RGBColor(31, 78, 121)
            day_para.paragraph_format.space_before = Pt(8)

            # Activities
            for act in day_data.get("activities", []):
                _bullet(doc, act)

            # Meals
            meals = day_data.get("meals", {})
            if meals:
                meals_para = doc.add_paragraph()
                meals_para.add_run("Meals: ").bold = True
                meals_para.add_run(
                    f"🍳 {meals.get('breakfast', '—')}  |  "
                    f"🥗 {meals.get('lunch', '—')}  |  "
                    f"🍽️ {meals.get('dinner', '—')}"
                )
                meals_para.paragraph_format.left_indent = Inches(0.25)

            # Cost
            cost_para = doc.add_paragraph()
            cost_para.add_run("Estimated Cost: ").bold = True
            cost_para.add_run(f"Rs.{_inr(day_cost)}")
            cost_para.paragraph_format.left_indent = Inches(0.25)
            cost_para.paragraph_format.space_after = Pt(6)

            _add_horizontal_rule(doc)

    # ======================================================
    # TRAVEL TIPS
    # ======================================================
    tips = itinerary.get("tips", [])
    if tips:
        _add_colored_heading(doc, "💡  Travel Tips & Recommendations", level=2)
        for tip in tips:
            _bullet(doc, tip)
        doc.add_paragraph()

    # ======================================================
    # FOOTER NOTE
    # ======================================================
    footer_para = doc.add_paragraph(
        "📌 This itinerary was autonomously generated by the AI Travel Agent. "
        "Prices and conditions are indicative. Always verify with service providers."
    )
    footer_para.paragraph_format.space_before = Pt(12)
    for run in footer_para.runs:
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)

    # ======================================================
    # SAVE FILE
    # ======================================================
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename  = f"itinerary_{user_id}_{timestamp}.docx"
    filepath  = os.path.join(output_dir, filename)
    doc.save(filepath)
    return os.path.abspath(filepath)
